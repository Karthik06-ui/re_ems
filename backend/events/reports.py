import os
import io
import zipfile
import subprocess
from datetime import datetime
from django.conf import settings
from django.utils import timezone
from django.contrib.auth import get_user_model
from django.db import transaction
from django.template.loader import render_to_string
from xhtml2pdf import pisa

from .models import Event, Registration, Team, TeamMember, EventAsset, Report, ReportVersion

User = get_user_model()



def ensure_directories():
    """Ensure media directories for templates, reports, and assets exist."""
    dirs = [
        os.path.join(settings.MEDIA_ROOT, 'templates'),
        os.path.join(settings.MEDIA_ROOT, 'generated_reports', 'docx'),
        os.path.join(settings.MEDIA_ROOT, 'generated_reports', 'pdf'),
        os.path.join(settings.MEDIA_ROOT, 'event_assets'),
    ]
    for d in dirs:
        os.makedirs(d, exist_ok=True)


def create_default_template(template_path):
    """
    Copy the pre-designed corporate report template from the repository templates folder,
    or programmatically generate a fallback template with the correct 9-table structure
    expected by the compiler if the source template is missing.
    """
    if os.path.exists(template_path):
        return

    # Try copying from repository templates first
    repo_template = os.path.join(settings.BASE_DIR, 'templates', 'event_report_template.docx')
    if os.path.exists(repo_template):
        import shutil
        try:
            shutil.copy2(repo_template, template_path)
            return
        except Exception as e:
            print(f"Error copying template from {repo_template}: {e}. Generating fallback.")

    # Programmatic fallback: Generate the exact 9-table structure to avoid IndexError
    import docx
    from docx.shared import Inches
    doc = docx.Document()
    
    # Page setup - 1 inch margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Table 0: Prepared By Metadata (2 rows, 4 columns)
    doc.add_heading("Prepared By Metadata", level=1)
    t0 = doc.add_table(rows=2, cols=4)
    t0.style = 'Table Grid'
    for r in t0.rows:
        for c in r.cells:
            c.text = "N/A"
    doc.add_paragraph("\n")

    # Table 1: Classification & Event Details (24 rows, 2 columns)
    doc.add_heading("Classification & Event Details", level=1)
    t1 = doc.add_table(rows=24, cols=2)
    t1.style = 'Table Grid'
    for r in t1.rows:
        for c in r.cells:
            c.text = "N/A"
    doc.add_paragraph("\n")

    # Table 2: Objective & Description (1 row, 1 column)
    doc.add_heading("Objective & Description", level=1)
    t2 = doc.add_table(rows=1, cols=1)
    t2.style = 'Table Grid'
    t2.rows[0].cells[0].text = "N/A"
    doc.add_paragraph("\n")

    # Table 3: Team Details / Participant List (2 rows, 6 columns)
    doc.add_heading("Team Details / Participant List", level=1)
    t3 = doc.add_table(rows=2, cols=6)
    t3.style = 'Table Grid'
    headers = ["S.No", "Name", "Roll Number", "Department", "Section/Year", "Role/Team"]
    for i, h in enumerate(headers):
        t3.rows[0].cells[i].text = h
    for c in t3.rows[1].cells:
        c.text = "N/A"
    doc.add_paragraph("\n")

    # Table 4: Event Summary & Highlights (1 row, 1 column)
    doc.add_heading("Event Summary & Highlights", level=1)
    t4 = doc.add_table(rows=1, cols=1)
    t4.style = 'Table Grid'
    t4.rows[0].cells[0].text = "N/A"
    doc.add_paragraph("\n")

    # Table 5: Outcomes & Key Takeaways (1 row, 1 column)
    doc.add_heading("Outcomes & Key Takeaways", level=1)
    t5 = doc.add_table(rows=1, cols=1)
    t5.style = 'Table Grid'
    t5.rows[0].cells[0].text = "N/A"
    doc.add_paragraph("\n")

    # Table 6: Photographs (2 rows, 2 columns)
    doc.add_heading("Photographs", level=1)
    t6 = doc.add_table(rows=2, cols=2)
    t6.style = 'Table Grid'
    for r in t6.rows:
        for c in r.cells:
            c.text = "N/A"
    doc.add_paragraph("\n")

    # Table 7: Supporting Documents Checklist (7 rows, 2 columns)
    doc.add_heading("Supporting Documents Checklist", level=1)
    t7 = doc.add_table(rows=7, cols=2)
    t7.style = 'Table Grid'
    checklist_labels = [
        "Document Category", "Status",
        "Invitation / Poster", "☐",
        "Certificates", "☐",
        "Newspaper / Social Media Proof", "☐",
        "Attendance / Participation List", "☐",
        "Budget / Expense Proof", "☐",
        "Other Docs", "☐"
    ]
    for idx in range(7):
        t7.rows[idx].cells[0].text = checklist_labels[idx * 2]
        t7.rows[idx].cells[1].text = checklist_labels[idx * 2 + 1]
    doc.add_paragraph("\n")

    # Table 8: Institutional Approval Sign-offs (2 rows, 3 columns)
    doc.add_heading("Institutional Approval Sign-offs", level=1)
    t8 = doc.add_table(rows=2, cols=3)
    t8.style = 'Table Grid'
    t8.rows[0].cells[0].text = "Prepared By:"
    t8.rows[0].cells[1].text = "Verified By:"
    t8.rows[0].cells[2].text = "Approved By:"
    for c in t8.rows[1].cells:
        c.text = "N/A"

    doc.save(template_path)






def set_cell_text(cell, text):
    """
    Set cell text by modifying the first paragraph and removing extra paragraphs.
    This preserves the cell and paragraph styling from the template.
    """
    if cell.paragraphs:
        p = cell.paragraphs[0]
        p.text = text
        # Remove any extra paragraphs in the cell to avoid unexpected spacing
        for extra_p in cell.paragraphs[1:]:
            p_element = extra_p._p
            p_element.getparent().remove(p_element)
    else:
        cell.text = text


@transaction.atomic
def build_event_report(event_id, requested_by=None):
    """
    Compile event report context, load the custom DOCX template,
    programmatically populate all standard metadata tables, insert participant lists,
    insert photographs, render checkmarks for supporting documents,
    convert to PDF, and log the version record.
    """
    import docx
    from docx.shared import Inches

    ensure_directories()
    
    event = Event.objects.get(id=event_id)
    report = getattr(event, 'report', None)
    
    if not report:
        raise ValueError("Report completion data has not been initialized for this event.")
        
    template_path = os.path.join(settings.MEDIA_ROOT, 'templates', 'event_report_template.docx')
    create_default_template(template_path)
    
    # Open the custom template programmatically using python-docx
    doc = docx.Document(template_path)
    
    # 1. Format text metadata
    start_str = event.start_time.strftime("%B %d, %Y %I:%M %p")
    end_str = event.end_time.strftime("%B %d, %Y %I:%M %p")
    date_range = f"From: {event.start_time.strftime('%Y-%m-%d %H:%M')}    To: {event.end_time.strftime('%Y-%m-%d %H:%M')}"
    
    # Calculate duration
    delta = event.end_time - event.start_time
    hours, remainder = divmod(delta.total_seconds(), 3600)
    minutes, _ = divmod(remainder, 60)
    duration = f"{int(hours)}h {int(minutes)}m" if hours > 0 else f"{int(minutes)}m"

    # Aggregates
    reg_count = event.registrations.exclude(status=Registration.Status.CANCELLED).count()
    att_count = event.registrations.filter(status=Registration.Status.CHECKED_IN).count()
    att_rate = round((att_count / reg_count * 100), 1) if reg_count > 0 else 0.0

    # Sponsors
    sponsors_list = [es.sponsor.name for es in event.event_sponsors.all()]
    sponsored_by = ", ".join(sponsors_list) if sponsors_list else "No External Sponsors"
    
    # --- Populating Table 0 (Prepared By Metadata) ---
    t0 = doc.tables[0]
    set_cell_text(t0.rows[0].cells[1], event.created_by_user.name if event.created_by_user else "Platform Admin")
    set_cell_text(t0.rows[0].cells[3], timezone.now().strftime("%Y-%m-%d"))
    set_cell_text(t0.rows[1].cells[1], "Organizer")
    set_cell_text(t0.rows[1].cells[3], "Final")

    # --- Populating Table 1 (Classification & Event Details) ---
    t1 = doc.tables[1]
    
    # Row 1: RiG Vertical
    set_cell_text(t1.rows[1].cells[1], event.rig_vertical or "Not Specified")
    # Row 2: Domain / Team
    set_cell_text(t1.rows[2].cells[1], event.domain_team or "Not Specified")
    # Row 3: Activity Type
    set_cell_text(t1.rows[3].cells[1], event.get_category_display())
    # Row 4: Name of Activity / Event / Workshop
    set_cell_text(t1.rows[4].cells[1], event.title)
    # Row 5: Event Type
    set_cell_text(t1.rows[5].cells[1], "Organized")
    # Row 6: Event ID
    set_cell_text(t1.rows[6].cells[1], str(event.id))
    # Row 8: Venue
    set_cell_text(t1.rows[8].cells[1], event.venue or "Online / Virtual")
    # Row 9: Mode
    set_cell_text(t1.rows[9].cells[1], event.get_type_display())
    # Row 10: Date
    set_cell_text(t1.rows[10].cells[1], date_range)
    # Row 11: Duration
    set_cell_text(t1.rows[11].cells[1], duration)
    # Row 13: No. of Participants / Competitors
    set_cell_text(t1.rows[13].cells[1], f"{att_count} / {reg_count} (Attended / Registered)")
    # Row 14: Organized By
    set_cell_text(t1.rows[14].cells[1], event.created_by_user.name if event.created_by_user else "Platform Admin")
    # Row 15: Coordinated By
    set_cell_text(t1.rows[15].cells[1], event.coordinated_by.name if event.coordinated_by else "Not Assigned")
    # Row 17: Prize Position
    set_cell_text(t1.rows[17].cells[1], report.prize_position or "Not Applicable")
    # Row 18: Prize Details
    set_cell_text(t1.rows[18].cells[1], report.prize_details or "Not Applicable")
    # Row 20: Sponsored By
    set_cell_text(t1.rows[20].cells[1], sponsored_by)
    # Row 21: Sponsored Amount
    set_cell_text(t1.rows[21].cells[1], f"Rs. {report.sponsored_amount:,.2f}")
    # Row 22: Amount Utilized
    set_cell_text(t1.rows[22].cells[1], f"Rs. {report.amount_utilized:,.2f}")
    # Row 23: Amount Returned
    set_cell_text(t1.rows[23].cells[1], f"Rs. {report.amount_returned:,.2f}")

    # --- Populating Table 2 (Objective & Brief Description) ---
    t2 = doc.tables[2]
    set_cell_text(t2.rows[0].cells[0], event.description or "No objective or description provided.")

    # --- Populating Table 3 (Team Details / Participant List) ---
    t3 = doc.tables[3]
    regs = event.registrations.filter(
        status__in=[Registration.Status.CONFIRMED, Registration.Status.CHECKED_IN]
    ).select_related('user', 'team')
    
    # Fill standard rows
    for idx, reg in enumerate(regs):
        if idx + 1 < len(t3.rows):
            row = t3.rows[idx + 1]
        else:
            row = t3.add_row()
        
        set_cell_text(row.cells[0], str(idx + 1))
        set_cell_text(row.cells[1], reg.user.name)
        set_cell_text(row.cells[2], reg.user.roll_number or "N/A")
        set_cell_text(row.cells[3], reg.user.department or "N/A")
        set_cell_text(row.cells[4], reg.user.year_of_study or "N/A")
        set_cell_text(row.cells[5], "Participant" if not reg.team else f"Team: {reg.team.name}")
        
    # Clear any unused pre-formatted rows in Table 3
    if len(regs) < 8:
        for idx in range(len(regs) + 1, len(t3.rows)):
            for cell in t3.rows[idx].cells:
                set_cell_text(cell, "")

    # --- Populating Table 4 (Event Summary & Highlights) ---
    t4 = doc.tables[4]
    set_cell_text(t4.rows[0].cells[0], report.summary or "Executive event summary not provided.")

    # --- Populating Table 5 (Outcomes & Key Takeaways) ---
    t5 = doc.tables[5]
    set_cell_text(t5.rows[0].cells[0], report.outcomes or "Outcomes and key takeaways not provided.")

    # --- Populating Table 6 (Photographs) ---
    t6 = doc.tables[6]
    photo_assets = event.assets.filter(category=EventAsset.AssetCategory.PHOTO)
    
    # Expand Table 6 if there are more than 4 photos
    if len(photo_assets) > 4:
        extra_rows_needed = ((len(photo_assets) - 1) // 2) + 1 - 2
        for _ in range(extra_rows_needed):
            t6.add_row()
            
    for idx, asset in enumerate(photo_assets):
        r_idx = idx // 2
        c_idx = idx % 2
        cell = t6.cell(r_idx, c_idx)
        # Clear existing text and paragraphs, then add picture
        if cell.paragraphs:
            p = cell.paragraphs[0]
            p.text = ""
            run = p.add_run()
            try:
                run.add_picture(asset.file.path, width=Inches(2.5))
                # Add caption as a new paragraph
                caption_p = cell.add_paragraph(f"Caption: {asset.name}")
                caption_p.alignment = 1
            except Exception as img_err:
                p.text = f"Error loading image: {asset.name}"
            # Remove any extra paragraphs beyond the first and the caption
            for extra_p in cell.paragraphs[2:]:
                p_element = extra_p._p
                p_element.getparent().remove(p_element)
        else:
            # Fallback
            cell.text = ""
            p = cell.add_paragraph()
            p.alignment = 1
            run = p.add_run()
            try:
                run.add_picture(asset.file.path, width=Inches(2.5))
                cell.add_paragraph(f"Caption: {asset.name}").alignment = 1
            except Exception as img_err:
                cell.text = f"Error loading image: {asset.name}"

    # --- Populating Table 7 (Supporting Documents Checklist) ---
    t7 = doc.tables[7]
    categories_checklist = {
        EventAsset.AssetCategory.INVITATION_POSTER: 1,
        EventAsset.AssetCategory.CERTIFICATE: 2,
        EventAsset.AssetCategory.MEDIA_PROOF: 3,
        EventAsset.AssetCategory.ATTENDANCE_LIST: 4,
        EventAsset.AssetCategory.BUDGET_PROOF: 5,
        EventAsset.AssetCategory.OTHER: 6,
    }
    
    for cat, row_idx in categories_checklist.items():
        has_file = event.assets.filter(category=cat).exists()
        set_cell_text(t7.rows[row_idx].cells[1], "☑" if has_file else "☐")

    # --- Appendix: Supporting Documents (embedded images + file references) ---
    # Non-photo assets grouped by category, appended after the checklist table
    supporting_assets = event.assets.exclude(category=EventAsset.AssetCategory.PHOTO)

    if supporting_assets.exists():
        IMAGE_EXTS = {'.jpg', '.jpeg', '.png', '.gif', '.webp', '.bmp', '.tiff'}
        CATEGORY_LABELS = {
            EventAsset.AssetCategory.INVITATION_POSTER: "Invitation / Poster",
            EventAsset.AssetCategory.CERTIFICATE: "Certificates / Achievement",
            EventAsset.AssetCategory.MEDIA_PROOF: "Newspaper / Media / Social-media Proof",
            EventAsset.AssetCategory.ATTENDANCE_LIST: "Attendance / Participation List",
            EventAsset.AssetCategory.BUDGET_PROOF: "Budget / Expense Proof",
            EventAsset.AssetCategory.OTHER: "Other Supporting Documents",
        }

        # Section heading paragraph (page-break first)
        doc.add_page_break()
        heading_para = doc.add_paragraph()
        heading_run = heading_para.add_run("Appendix: Supporting Documents")
        heading_run.bold = True
        heading_run.font.size = docx.shared.Pt(14)
        heading_para.paragraph_format.space_after = docx.shared.Pt(6)

        for asset in supporting_assets:
            # Sub-heading for category
            sub_para = doc.add_paragraph()
            sub_run = sub_para.add_run(CATEGORY_LABELS.get(asset.category, asset.category.replace("_", " ").title()))
            sub_run.bold = True
            sub_run.font.size = docx.shared.Pt(11)

            # Asset name line
            name_para = doc.add_paragraph(f"File: {asset.name}")
            name_para.paragraph_format.space_before = docx.shared.Pt(2)
            name_para.paragraph_format.space_after = docx.shared.Pt(4)

            file_ext = os.path.splitext(asset.file.name)[1].lower()
            if file_ext in IMAGE_EXTS and os.path.exists(asset.file.path):
                try:
                    img_para = doc.add_paragraph()
                    img_para.alignment = 1  # centre
                    img_run = img_para.add_run()
                    img_run.add_picture(asset.file.path, width=Inches(5.5))
                except Exception:
                    doc.add_paragraph("[Image could not be embedded — see ZIP package]")
            else:
                doc.add_paragraph(f"(Non-image file — available in the ZIP download package: {os.path.basename(asset.file.name)})")

            # Separator
            doc.add_paragraph()

    # Version tracking
    latest_version = ReportVersion.objects.filter(event=event).order_by('-version_number').first()
    new_version_num = (latest_version.version_number + 1) if latest_version else 1

    # Output filenames
    docx_filename = f"event_{event.id}_v{new_version_num}.docx"
    pdf_filename = f"event_{event.id}_v{new_version_num}.pdf"
    
    docx_dir = os.path.join(settings.MEDIA_ROOT, 'generated_reports', 'docx')
    pdf_dir = os.path.join(settings.MEDIA_ROOT, 'generated_reports', 'pdf')
    
    docx_output_path = os.path.join(docx_dir, docx_filename)
    pdf_output_path = os.path.join(pdf_dir, pdf_filename)

    # Save DOCX
    doc.save(docx_output_path)

    # Generate PDF using WeasyPrint (HTML to PDF pipeline)
    regs = event.registrations.filter(
        status__in=[Registration.Status.CONFIRMED, Registration.Status.CHECKED_IN]
    ).select_related('user', 'team')
    
    event_assets = event.assets.all()
    
    assets_dict = {
        'invitation_poster': False,
        'certificate': False,
        'media_proof': False,
        'attendance_list': False,
        'budget_proof': False,
        'other': False
    }
    for asset in event_assets:
        if asset.category == EventAsset.AssetCategory.PHOTO:
            pass # Handled differently or ignored in checklist
        else:
            assets_dict[asset.category] = True

    html_string = render_to_string('events/report_template.html', {
        'event': event,
        'report': report,
        'registrations': regs,
        'assets': event_assets,
        'assets_dict': assets_dict,
        'generated_by': requested_by,
    })
    
    try:
        with open(pdf_output_path, "w+b") as out_pdf:
            pisa_status = pisa.CreatePDF(html_string, dest=out_pdf)
        if pisa_status.err:
            print(f"xhtml2pdf PDF compilation failed with errors")
    except Exception as e:
        import traceback
        traceback.print_exc()
        print(f"xhtml2pdf PDF compilation failed: {e}")

    # Update previous active versions to inactive
    ReportVersion.objects.filter(event=event).update(is_active=False)

    # Save to db
    version_instance = ReportVersion.objects.create(
        event=event,
        version_number=new_version_num,
        docx_file=os.path.relpath(docx_output_path, settings.MEDIA_ROOT).replace("\\", "/"),
        pdf_file=os.path.relpath(pdf_output_path, settings.MEDIA_ROOT).replace("\\", "/"),
        generated_by=requested_by,
        is_active=True
    )

    return version_instance


def compile_report_zip(event_id):
    """
    Compile the active report version files (DOCX, PDF) and all event assets
    into a ZIP package in-memory. Returns a BytesIO stream.
    """
    event = Event.objects.get(id=event_id)
    active_version = ReportVersion.objects.filter(event=event, is_active=True).first()
    
    if not active_version:
        raise ValueError("No compiled report version found. Please generate the report before downloading the ZIP archive.")

    zip_buffer = io.BytesIO()
    
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        # 1. Add DOCX
        if active_version.docx_file and os.path.exists(active_version.docx_file.path):
            zip_file.write(
                active_version.docx_file.path, 
                arcname=f"Event_Report_v{active_version.version_number}.docx"
            )
            
        # 2. Add PDF
        if active_version.pdf_file and os.path.exists(active_version.pdf_file.path):
            zip_file.write(
                active_version.pdf_file.path, 
                arcname=f"Event_Report_v{active_version.version_number}.pdf"
            )

        # 3. Add Event Assets grouped by categories
        event_assets = event.assets.all()
        for asset in event_assets:
            if asset.file and os.path.exists(asset.file.path):
                # Clean name and group in folder structure
                folder_name = "Photographs" if asset.category == EventAsset.AssetCategory.PHOTO else "Supporting_Documents"
                category_folder = asset.category.replace("_", " ").title().replace(" ", "_")
                
                if asset.category == EventAsset.AssetCategory.PHOTO:
                    archive_path = os.path.join(folder_name, asset.name)
                else:
                    archive_path = os.path.join(folder_name, category_folder, asset.name)
                    
                zip_file.write(asset.file.path, arcname=archive_path)

    zip_buffer.seek(0)
    return zip_buffer
